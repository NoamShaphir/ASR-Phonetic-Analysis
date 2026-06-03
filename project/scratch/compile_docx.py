import os
import re
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_shading(cell, color_hex):
    """Applies background color to a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding (in twips)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3"):
    """Applies clean horizontal borders and removes vertical ones."""
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
            f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="{color}"/>'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
            f'  <w:left w:val="none"/>'
            f'  <w:right w:val="none"/>'
            f'  <w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def add_styled_run(paragraph, text, is_bold=False, is_italic=False, is_code=False, font_size=11, color_rgb=None):
    """Helper to add runs with inline formatting (bold, italic, code)."""
    run = paragraph.add_run(text)
    run.font.name = 'Consolas' if is_code else 'Arial'
    run.font.size = Pt(font_size)
    run.bold = is_bold
    run.italic = is_italic
    if color_rgb:
        run.font.color.rgb = color_rgb
    return run

def parse_and_add_inline_formatting(paragraph, md_text, font_size=11, base_color=None):
    """Parses markdown inline bold (**), italic (*), and code (`) and adds runs to paragraph."""
    # Pattern to match bold, italic, code, and plain text
    # e.g., **bold**, *italic*, `code`, plain text
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|[^\*`]+)')
    matches = pattern.findall(md_text)
    
    for part in matches:
        if part.startswith('**') and part.endswith('**'):
            content = part[2:-2]
            add_styled_run(paragraph, content, is_bold=True, font_size=font_size, color_rgb=base_color)
        elif part.startswith('*') and part.endswith('*'):
            content = part[1:-1]
            add_styled_run(paragraph, content, is_italic=True, font_size=font_size, color_rgb=base_color)
        elif part.startswith('`') and part.endswith('`'):
            content = part[1:-1]
            add_styled_run(paragraph, content, is_code=True, font_size=font_size, color_rgb=RGBColor(216, 27, 96)) # Coral-pink for code
        else:
            # Check for LaTeX style equations and make them italic or math-like
            if part.strip().startswith('$') and part.strip().endswith('$'):
                content = part.strip()[1:-1]
                add_styled_run(paragraph, content, is_italic=True, font_size=font_size, color_rgb=RGBColor(0, 102, 204))
            else:
                add_styled_run(paragraph, part, font_size=font_size, color_rgb=base_color)

def compile_markdown_to_docx(md_path, docx_path):
    print(f"Reading markdown from: {md_path}")
    if not os.path.exists(md_path):
        raise FileNotFoundError(f"Markdown file not found: {md_path}")
        
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
        
    doc = Document()
    
    # Page setup - 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles Setup (Arial base)
    primary_color = RGBColor(43, 45, 66)      # Dark Slate Blue #2B2D42
    secondary_color = RGBColor(141, 153, 174) # Cool Grey #8D99AE
    dark_grey = RGBColor(51, 51, 51)          # #333333
    
    lines = md_content.split('\n')
    
    in_table = False
    table_headers = []
    table_rows = []
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Check for mermaid or code block start - skip them or format them as callout blocks
        if line.startswith('```'):
            code_block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_block.append(lines[i])
                i += 1
            
            # Write a callout block for the code
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.right_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            
            # Soft grey background and border
            p_format = p._element.get_or_add_pPr()
            pbdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="8" w:color="8D99AE"/></w:pBdr>')
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8F9FA"/>')
            p_format.append(pbdr)
            p_format.append(shd)
            
            code_text = "\n".join(code_block)
            add_styled_run(p, code_text, is_code=True, font_size=9.5, color_rgb=dark_grey)
            i += 1
            continue
            
        # Parse Markdown Tables
        if line.startswith('|'):
            in_table = True
            cells = [c.strip() for c in line.split('|')[1:-1]]
            
            # If it's the separator row, skip it
            if all(re.match(r'^:?-+:?$', c) for c in cells):
                i += 1
                continue
                
            if not table_headers:
                table_headers = cells
            else:
                table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table:
                # We finished a table, let's render it in docx!
                if table_headers:
                    cols_count = len(table_headers)
                    table = doc.add_table(rows=1, cols=cols_count)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    # Set borders and header background
                    set_table_borders(table, color="D3D3D3")
                    
                    # Add headers
                    hdr_cells = table.rows[0].cells
                    for col_idx, cell_text in enumerate(table_headers):
                        hdr_cells[col_idx].text = ""
                        p = hdr_cells[col_idx].paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p.paragraph_format.space_after = Pt(4)
                        p.paragraph_format.space_before = Pt(4)
                        set_cell_shading(hdr_cells[col_idx], "2B2D42") # Primary Dark Slate Blue
                        set_cell_margins(hdr_cells[col_idx], top=120, bottom=120, left=140, right=140)
                        
                        # Add header text run (Bold, White)
                        parse_and_add_inline_formatting(p, cell_text, font_size=10, base_color=RGBColor(255, 255, 255))
                        
                    # Add rows
                    for r_idx, row_data in enumerate(table_rows):
                        # Pad row_data if short
                        if len(row_data) < cols_count:
                            row_data += [""] * (cols_count - len(row_data))
                        
                        row = table.add_row()
                        row_cells = row.cells
                        
                        bg_color = "F8F9FA" if r_idx % 2 == 1 else "FFFFFF" # Zebra striping
                        
                        for col_idx, cell_text in enumerate(row_data[:cols_count]):
                            row_cells[col_idx].text = ""
                            p = row_cells[col_idx].paragraphs[0]
                            p.paragraph_format.space_after = Pt(3)
                            p.paragraph_format.space_before = Pt(3)
                            
                            set_cell_shading(row_cells[col_idx], bg_color)
                            set_cell_margins(row_cells[col_idx], top=90, bottom=90, left=140, right=140)
                            row_cells[col_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            
                            parse_and_add_inline_formatting(p, cell_text, font_size=9.5, base_color=dark_grey)
                    
                    # Space after table
                    doc.add_paragraph().paragraph_format.space_before = Pt(8)
                    
                in_table = False
                table_headers = []
                table_rows = []
                
        # Parse Headings
        if line.startswith('#'):
            h_match = re.match(r'^(#+)\s*(.*)$', line)
            if h_match:
                level = len(h_match.group(1))
                title = h_match.group(2).strip()
                
                if level == 1:
                    # Document Main Title
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(18)
                    p.paragraph_format.space_after = Pt(12)
                    parse_and_add_inline_formatting(p, title, font_size=24, base_color=primary_color)
                    p.runs[0].bold = True
                    
                    # Add subtle divider line under title
                    p_format = p._element.get_or_add_pPr()
                    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="18" w:space="6" w:color="2B2D42"/></w:pBdr>')
                    p_format.append(pBdr)
                    
                elif level == 2:
                    # Heading 2
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(18)
                    p.paragraph_format.space_after = Pt(8)
                    p.paragraph_format.keep_with_next = True
                    parse_and_add_inline_formatting(p, title, font_size=15, base_color=primary_color)
                    p.runs[0].bold = True
                    
                    # Add a divider for H2 as well
                    p_format = p._element.get_or_add_pPr()
                    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="3" w:color="8D99AE"/></w:pBdr>')
                    p_format.append(pBdr)
                    
                elif level == 3:
                    # Heading 3
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(14)
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.keep_with_next = True
                    parse_and_add_inline_formatting(p, title, font_size=12, base_color=primary_color)
                    p.runs[0].bold = True
                    
                else:
                    # Heading 4+
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(4)
                    p.paragraph_format.keep_with_next = True
                    parse_and_add_inline_formatting(p, title, font_size=11, base_color=dark_grey)
                    p.runs[0].bold = True
                
            i += 1
            continue
            
        # Parse Lists
        list_match = re.match(r'^([\-\*\+])\s*(.*)$', line)
        num_list_match = re.match(r'^(\d+)\.\s*(.*)$', line)
        
        if list_match:
            bullet_content = list_match.group(2).strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.left_indent = Inches(0.4)
            parse_and_add_inline_formatting(p, bullet_content, font_size=10.5, base_color=dark_grey)
            i += 1
            continue
        elif num_list_match:
            num_content = num_list_match.group(2).strip()
            num_str = num_list_match.group(1)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.left_indent = Inches(0.4)
            
            # Add custom formatted number to allow styled list numbers
            run_num = add_styled_run(p, f"{num_str}.  ", is_bold=True, font_size=10.5, color_rgb=primary_color)
            parse_and_add_inline_formatting(p, num_content, font_size=10.5, base_color=dark_grey)
            i += 1
            continue
            
        # Parse Horizontal Divider
        if line == '---' or line == '***':
            # Create a stylized paragraph divider
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run("♦   ♦   ♦")
            run.font.size = Pt(8)
            run.font.color.rgb = secondary_color
            i += 1
            continue
            
        # Parse Standard Paragraphs
        if line:
            # Check for GitHub-style Alert boxes
            # e.g., > [!NOTE] or > [!IMPORTANT]
            if line.startswith('>'):
                alert_type = "NOTE"
                alert_text = line[1:].strip()
                
                # Check if next lines also belong to the blockquote
                blockquote_lines = [alert_text]
                i += 1
                while i < len(lines) and lines[i].strip().startswith('>'):
                    blockquote_lines.append(lines[i].strip()[1:].strip())
                    i += 1
                
                # Compile blockquote text
                full_bq_text = " ".join(blockquote_lines)
                
                # Detect alert tag
                tag_match = re.match(r'^\[!(IMPORTANT|NOTE|WARNING|TIP|CAUTION)\]\s*(.*)$', full_bq_text)
                if tag_match:
                    alert_type = tag_match.group(1)
                    alert_text = tag_match.group(2)
                else:
                    alert_text = full_bq_text
                
                # Setup Alert Box Colors
                # Default Note (Blue)
                border_color = "0066CC"
                shd_color = "F0F7FF"
                alert_rgb = RGBColor(0, 102, 204)
                
                if alert_type == "IMPORTANT":
                    border_color = "D81B60" # Dark Pink/Red
                    shd_color = "FFF0F5"
                    alert_rgb = RGBColor(216, 27, 96)
                elif alert_type == "WARNING" or alert_type == "CAUTION":
                    border_color = "FF9900" # Orange
                    shd_color = "FFF9E6"
                    alert_rgb = RGBColor(255, 153, 0)
                elif alert_type == "TIP":
                    border_color = "2A9D8F" # Greenish Teal
                    shd_color = "E6F4F2"
                    alert_rgb = RGBColor(42, 157, 143)
                
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.right_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(8)
                
                p_format = p._element.get_or_add_pPr()
                pbdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="36" w:space="12" w:color="{border_color}"/></w:pBdr>')
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shd_color}"/>')
                p_format.append(pbdr)
                p_format.append(shd)
                
                # Add alert title
                run_tag = add_styled_run(p, f"★  {alert_type}: ", is_bold=True, font_size=10, color_rgb=alert_rgb)
                parse_and_add_inline_formatting(p, alert_text, font_size=10, base_color=dark_grey)
                
                continue # Already incremented i
                
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            
            parse_and_add_inline_formatting(p, line, font_size=10.5, base_color=dark_grey)
            
        i += 1
        
    print(f"Saving compiled Word document to: {docx_path}")
    try:
        doc.save(docx_path)
        print("[+] Success!")
        return True
    except PermissionError:
        print(f"[-] Permission Error: Could not save to {docx_path}. Is the file open in MS Word?")
        # Try a fallback name
        fallback_path = docx_path.replace(".docx", "_New.docx")
        print(f"Trying to save to fallback path instead: {fallback_path}")
        doc.save(fallback_path)
        print("[+] Saved successfully to fallback path!")
        return fallback_path

if __name__ == '__main__':
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_dir = os.path.dirname(script_dir)
    
    # Master Report paths
    master_md = os.path.join(project_dir, "docs", "ASR_Phonetic_Analysis_Master_Report.md")
    master_docx = os.path.join(project_dir, "docs", "ASR_Phonetic_Analysis_Master_Report.docx")
    
    # New Test Dataset Report paths
    test_md = os.path.join(project_dir, "docs", "New_Test_Dataset_Evaluation_Report.md")
    test_docx = os.path.join(project_dir, "docs", "New_Test_Dataset_Evaluation_Report.docx")
    
    print("=== DOCX COMPILER START ===")
    
    print("Compiling ASR Phonetic Analysis Master Report...")
    compile_markdown_to_docx(master_md, master_docx)
    
    print("\nCompiling New Test Dataset Evaluation Report...")
    compile_markdown_to_docx(test_md, test_docx)
    
    print("=== DOCX COMPILER END ===")
